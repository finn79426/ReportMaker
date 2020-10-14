from docx import Document
from copy import copy
from openpyxl import load_workbook
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from PIL import Image
import os

INPUTXLSX = "./vulns.xlsx"
OUTPUTDOCX = "./output.docx"


def copyFormInWord(times):
    template = Document("./template.docx")
    template_form = template.tables[0]
    for i in range(1, times):
        paragraph = template.add_paragraph()
        # WordprocessingML: _tbl, _p
        new_tbl = copy(template_form._tbl)
        paragraph._p.addnext(new_tbl)
    template.save(OUTPUTDOCX)


def readRow(sheet, row):
    result = []
    for target_row in sheet.iter_rows(min_row=row, max_row=row):
        for cell in target_row:
            result.append(cell.value)
    return result


def writeWord(Data, nTable):
    report = Document(OUTPUTDOCX)
    currentTable = report.tables[nTable-1]  # .tables start at 0

    # Init
    styleWord = report.styles['Normal']
    styleWord._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')
    fontWord = styleWord.font
    fontWord.size = Pt(12)
    fontWord.name = "微軟正黑體"
    fontWord.color.rgb = RGBColor(0, 0, 0)
    currentTable.autofit = False    # 表格欄寬固定不變，文字配合表格進行換行
    paragraph = report.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0   # 單行間距

    # Text
    #print(Data)
    currentTable.cell(0, 1).text = Data[2]  # Vulnerability
    currentTable.cell(0, 3).text = Data[0]  # Site
    currentTable.cell(1, 1).text = Data[1]  # URL(s)
    currentTable.cell(2, 1).text = Data[5]  # Risk
    currentTable.cell(3, 1).text = Data[3]  # Details
    currentTable.cell(4, 1).text = Data[4]  # Remedy

    print("\t", "Vulnerability:", currentTable.cell(0, 1).text)
    print("\t", "Site:", currentTable.cell(0, 3).text)
    print("\t", "Risk:", currentTable.cell(2, 1).text)

    # Picture
    dirName = Data[6]
    dirLocation = './proofs/' + str(dirName)
    if os.path.isdir(dirLocation) and dirName != None and len(os.listdir(dirLocation)) != 0:
        ls = os.listdir(dirLocation)
        for fname in ls:
            im = Image.open(str(dirLocation + '/' + fname))
            im.thumbnail((1028, 1028), Image.LANCZOS)
            picWidth, picHeight = im.size
            picRow = currentTable.rows[5]
            picColumn = picRow.cells[2]
            picParagraph = picColumn.paragraphs[0]
            add_pic = picParagraph.add_run()
            add_pic.add_picture(str(dirLocation+'/'+fname), width=Cm(
                ((picWidth*2.54)/96)/2), height=Cm(((picHeight*2.54)/96)/2))
            print("\t", str(dirLocation+'/'+fname), im.size)
    else:
        pass
    report.save(OUTPUTDOCX)


if __name__ == "__main__":
    sheet = load_workbook(INPUTXLSX).active
    vulncount = sheet.max_row - 1
    print("\033[92m", "Total Vulnerability:", vulncount, "\033[0m")
    copyFormInWord(vulncount)
    for i in range(1, vulncount+1):     # 第一列是標題，跳過不計算
        print("\033[94m", "[*] Processing Row", i+1, "\033[0m")
        formData = readRow(sheet, i+1)  # 第一列是標題，跳過不計算
        writeWord(formData, i)
